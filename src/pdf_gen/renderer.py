"""
pdf_gen/renderer.py — Render tailored resumes and cover letters as PDFs.

Strategy:
  1. Copy the lane's base .docx template
  2. Surgically replace the writable sections with tailored content
  3. Convert .docx → .pdf via LibreOffice (primary) or docx2pdf (macOS fallback)

Resume paragraph map (base_resume.docx — do not change without re-inspecting):
  [3]  Tagline          — run[0] is bold tagline text; run[1] is a trailing space
  [6]  Summary          — plain Times New Roman 10pt, 2+ runs
  [9]  (spacer before the skills table)
  [12] Role 0 bullet 0  — Primary role (intro/first bullet)
  [13] (empty spacer — leave alone)
  [14] Role 0 bullet 1
  [15] Role 0 bullet 2
  [16] Role 0 bullet 3
  [17] Role 0 bullet 4
  [20] Role 1 bullet 0  — Secondary role
  [23] Role 2 bullet 0  — Tertiary role

Skills table (Table 0):  3 rows × 3 cols
  Each cell has 2 runs:
    run[0]  ' ✓'  Arial Unicode MS 10pt  ← KEEP UNTOUCHED
    run[1]  ' Skill Name '  Times New Roman 10pt  ← REPLACE TEXT ONLY
"""

import os
import re
import shutil
import subprocess
import sys
from copy import deepcopy
from hashlib import blake2b
from pathlib import Path
from typing import Optional

import structlog
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

log = structlog.get_logger()

_ROOT = Path(__file__).resolve().parent.parent.parent

# ── Fixed indices in base_resume.docx ────────────────────────────────────────
_TAGLINE_IDX = 3
_SUMMARY_IDX = 6

# Summary cap — hard backstop so a runaway LLM response doesn't overflow the summary block.
_SUMMARY_MAX_CHARS = 420   # ~3 full sentences; sentence-boundary truncation only

# Bullet backstop — the LLM is instructed to stay under 150 chars, but if it doesn't,
# truncate at the last complete word so text never wraps onto a second line.
_BULLET_MAX_CHARS = 160

# bullet paragraph indices per role (index → list of para indices)
_ROLE_BULLET_IDXS: dict[int, list[int]] = {
    0: [12, 14, 15, 16, 17],   # Primary role (para 13 is an empty spacer — skip)
    1: [20],                    # Secondary role
    2: [23],                    # Tertiary role
    3: [26],                    # Fourth role (1-line summary only)
}

# LibreOffice binary locations to try in order
_LO_CANDIDATES = [
    os.getenv("LIBREOFFICE_PATH", ""),
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "soffice",
    "libreoffice",
]


# ── Public API ────────────────────────────────────────────────────────────────


def render_resume(
    tailored_resume: dict,
    lane: dict,
    job: dict,
    date_str: str,
    output_dir: Path,
) -> tuple[Optional[Path], Path]:
    """
    Fill the lane's base resume template with tailored content and convert to PDF.

    Returns (pdf_path or None, docx_path) — DOCX is always written to output_dir.
    The DOCX is the editable intermediate; the PDF is the converted final.
    Content is semantically identical because the PDF is generated from the
    same DOCX.

    If PDF conversion is unavailable (no LibreOffice / docx2pdf installed), the
    first element is None to give callers an unambiguous "no PDF" signal — they
    can then filter the None, log it, or fall back to a DOCX-only message
    without falsely claiming a PDF exists.
    """
    template_path = _ROOT / lane["template"]
    if not template_path.exists():
        raise FileNotFoundError(f"Resume template not found: {template_path}")

    output_dir.mkdir(parents=True, exist_ok=True)
    base = _compose_output_base(job, "Resume")
    docx_path = output_dir / f"{base}.docx"
    pdf_path  = output_dir / f"{base}.pdf"

    shutil.copy2(template_path, docx_path)
    docx_path.chmod(0o644)  # template may be read-only; we need to write the filled copy
    _fill_resume_template(docx_path, tailored_resume)
    _docx_to_pdf(docx_path, output_dir)

    if pdf_path.exists():
        log.info("renderer.resume_documents", pdf=str(pdf_path), docx=str(docx_path))
        return pdf_path, docx_path

    log.warning("renderer.pdf_unavailable_docx_only", docx=str(docx_path))
    return None, docx_path


def render_cover_letter(
    cover_letter: dict,
    job: dict,
    date_str: str,
    output_dir: Path,
) -> tuple[Optional[Path], Path]:
    """
    Build a clean cover letter DOCX from the paragraph list and convert to PDF.

    Returns (pdf_path or None, docx_path) — DOCX is always written to output_dir.
    The DOCX is the editable intermediate; the PDF is the converted final.

    If PDF conversion is unavailable, the first element is None — see
    render_resume() docstring for rationale.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    base = _compose_output_base(job, "Cover_Letter")
    docx_path = output_dir / f"{base}.docx"
    pdf_path  = output_dir / f"{base}.pdf"

    _create_cover_letter_docx(docx_path, cover_letter, job)
    _docx_to_pdf(docx_path, output_dir)

    if pdf_path.exists():
        log.info("renderer.cover_letter_documents", pdf=str(pdf_path), docx=str(docx_path))
        return pdf_path, docx_path

    log.warning("renderer.pdf_unavailable_docx_only", docx=str(docx_path))
    return None, docx_path


# ── Resume template filling ───────────────────────────────────────────────────


def _fill_resume_template(docx_path: Path, tailored: dict) -> None:
    """
    Open the copied template and surgically replace writable sections.
    All untouched paragraphs (name, contact, role headers, education, Kitchen Manager)
    are left exactly as-is — including their run-level formatting.
    """
    doc = Document(str(docx_path))
    paras = doc.paragraphs  # snapshot — object references remain valid after insertions

    # 1. Tagline — update run[0] text in place; run[1] is the trailing space, leave it
    tagline = tailored.get("tagline", "")
    if tagline and len(paras) > _TAGLINE_IDX:
        para = paras[_TAGLINE_IDX]
        if para.runs:
            para.runs[0].text = f" {tagline}"
            # Clear any additional text runs (run[1] is a spacer — keep, clear others)
            for run in para.runs[2:]:
                run.text = ""

    # 2. Summary — replace all runs with a single formatted run (capped for 1-page fit)
    summary = tailored.get("summary", "")
    if summary and len(paras) > _SUMMARY_IDX:
        _replace_para_runs(paras[_SUMMARY_IDX], _fit(summary, _SUMMARY_MAX_CHARS))

    # 3. Skills table — update run[1] (the text run) only; run[0] is the '✓' glyph
    skills = tailored.get("skills", [])
    if skills and doc.tables:
        skill_idx = 0
        for row in doc.tables[0].rows:
            for cell in row.cells:
                if skill_idx >= len(skills):
                    break
                cell_para = cell.paragraphs[0]
                if len(cell_para.runs) >= 2:
                    # run[0] = ' ✓' in Arial Unicode MS — keep intact
                    # run[1] = ' Skill Name ' in Times New Roman — replace text only
                    cell_para.runs[1].text = f" {skills[skill_idx]} "
                elif cell_para.runs:
                    # Unexpected structure — safe fallback
                    cell_para.runs[-1].text = f" {skills[skill_idx]} "
                skill_idx += 1

    # 4. Role bullets — replace existing slots; truncate if tailor returned too many,
    #    clear if tailor returned too few. Never insert paragraphs (would break layout).
    for role_dict in tailored.get("roles", []):
        role_idx = role_dict.get("index")
        if role_idx not in _ROLE_BULLET_IDXS:
            continue

        slots       = _ROLE_BULLET_IDXS[role_idx]
        new_bullets = role_dict.get("bullets", [])[:len(slots)]  # truncate to slot count

        for slot_i, para_idx in enumerate(slots):
            if para_idx >= len(paras):
                continue
            if slot_i < len(new_bullets):
                _replace_para_runs(paras[para_idx], _fit_bullet(new_bullets[slot_i]))
            else:
                _hide_bullet_para(paras[para_idx])  # clear unused slots — strip bullet glyph too

    doc.save(str(docx_path))
    log.info("renderer.resume_docx_saved", path=str(docx_path))


# ── Cover letter document ─────────────────────────────────────────────────────


def _create_cover_letter_docx(
    docx_path: Path,
    cover_letter: dict,
    job: dict,
) -> None:
    """
    Create a clean Calibri 11pt cover letter DOCX from the paragraph list.
    1" side margins, 0.75" top/bottom, 8pt space after each paragraph.
    """
    doc = Document()

    # Page margins: 1" sides, 0.75" top/bottom
    for section in doc.sections:
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # Default Normal style — Calibri 11pt
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Simple header: name + contact info
    # These come from environment or config — never hardcode personal data
    applicant_name = os.getenv("APPLICANT_NAME", "Your Name")
    applicant_contact = os.getenv("APPLICANT_CONTACT", "City, ST  ·  email@example.com  ·  (555) 000-0000")
    _add_cl_run(doc.add_paragraph(), applicant_name, bold=True, size=Pt(11))
    addr_para = doc.add_paragraph()
    _add_cl_run(addr_para, applicant_contact, size=Pt(10))
    addr_para.paragraph_format.space_after = Pt(14)

    # Body paragraphs
    paragraphs = cover_letter.get("paragraphs", [])
    for i, text in enumerate(paragraphs):
        p = doc.add_paragraph()
        _add_cl_run(p, text, size=Pt(11))
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Pt(0)

    doc.save(str(docx_path))
    log.info("renderer.cl_docx_saved", path=str(docx_path))


def _add_cl_run(para, text: str, bold: bool = False, size=None) -> None:
    """Add a single run to a paragraph with explicit Calibri formatting."""
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.bold = bold
    if size is not None:
        run.font.size = size


# ── DOCX → PDF conversion ─────────────────────────────────────────────────────


def _docx_to_pdf(docx_path: Path, output_dir: Path) -> None:
    """
    Convert docx → PDF.

    Tries in order:
      1. LibreOffice CLI (production path — works on Linux/macOS if installed)
      2. docx2pdf Python library (macOS with Microsoft Word installed)

    If neither is available, logs a clear install instruction and returns.
    The caller falls back to returning the .docx.
    """
    # Try LibreOffice
    lo_bin = _find_libreoffice()
    if lo_bin:
        _convert_with_libreoffice(lo_bin, docx_path, output_dir)
        return

    # Try docx2pdf (macOS with Word)
    if sys.platform == "darwin":
        try:
            from docx2pdf import convert
            convert(str(docx_path), str(docx_path.with_suffix(".pdf")))
            log.debug("renderer.docx2pdf_ok", path=str(docx_path))
            return
        except Exception as e:
            log.debug("renderer.docx2pdf_failed", error=str(e))

    log.warning(
        "renderer.no_pdf_converter",
        hint=(
            "Install LibreOffice to enable PDF generation.\n"
            "  macOS: download from https://www.libreoffice.org/download/download/\n"
            "  Linux: sudo apt install libreoffice-writer\n"
            "  Docker: see deploy/Dockerfile\n"
            "DOCX files are still generated and usable."
        ),
    )


def _find_libreoffice() -> str | None:
    """Return the path to a working LibreOffice binary, or None."""
    for candidate in _LO_CANDIDATES:
        if not candidate:
            continue
        try:
            result = subprocess.run(
                [candidate, "--version"],
                capture_output=True, text=True, timeout=5,
            )
            if result.returncode == 0:
                return candidate
        except (FileNotFoundError, subprocess.TimeoutExpired, PermissionError):
            continue
    return None


def _convert_with_libreoffice(lo_bin: str, docx_path: Path, output_dir: Path) -> None:
    cmd = [
        lo_bin, "--headless",
        "--convert-to", "pdf",
        "--outdir", str(output_dir),
        str(docx_path),
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        if result.returncode != 0:
            log.error("renderer.libreoffice_error",
                      stderr=result.stderr[:500], returncode=result.returncode)
        else:
            log.debug("renderer.libreoffice_ok", stdout=result.stdout[:200])
    except subprocess.TimeoutExpired:
        log.error("renderer.libreoffice_timeout")


# ── Paragraph helpers ─────────────────────────────────────────────────────────


def _fit_bullet(text: str) -> str:
    """
    Backstop for bullet text: truncate at the last complete word boundary if the bullet
    exceeds _BULLET_MAX_CHARS. The LLM is instructed to stay under 150 chars; this 160-char
    limit is a safety net only. Never cuts mid-word.
    """
    if len(text) <= _BULLET_MAX_CHARS:
        return text
    log.warning("renderer.bullet_truncated", original_len=len(text), preview=text[:60])
    window = text[:_BULLET_MAX_CHARS]
    last_space = window.rfind(" ")
    return window[:last_space] if last_space > 0 else window


def _fit(text: str, max_chars: int) -> str:
    """
    Truncate text to at most max_chars, cutting only at a sentence boundary ('. ').
    If no sentence boundary exists within the window, returns the full text untruncated
    rather than cutting mid-sentence or mid-phrase.
    Logs a warning if truncation occurs so LLM prompts can be tightened upstream.
    """
    if len(text) <= max_chars:
        return text

    log.warning("renderer.text_truncated", original_len=len(text), max=max_chars,
                preview=text[:60])

    window = text[:max_chars]
    last_period = window.rfind(". ")
    if last_period > 0:
        return window[:last_period + 1]

    # No clean sentence boundary — return full text rather than cutting mid-phrase
    return text


def _hide_bullet_para(para) -> None:
    """
    Clear an unused bullet slot so it is truly invisible:
      1. Remove all w:r (run) elements so no text is rendered.
      2. Strip w:numPr from w:pPr so the bullet glyph/indent also disappears.
    """
    p_elem = para._element
    for r_elem in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r_elem)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)


def _replace_para_runs(para, new_text: str) -> None:
    """
    Clear all w:r elements in a paragraph and replace with a single run.
    Preserves the formatting (bold, font name, size) from the first existing run.
    Paragraph-level properties (spacing, borders, indent) are never touched.
    """
    # Snapshot formatting before we destroy anything
    bold = None
    font_name = "Times New Roman"
    font_size = Pt(10)

    if para.runs:
        r0 = para.runs[0]
        bold      = r0.bold
        font_name = r0.font.name or "Times New Roman"
        font_size = r0.font.size or Pt(10)

    # Remove every w:r from the paragraph's XML element
    p_elem = para._element
    for r_elem in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r_elem)

    if not new_text:
        return

    run = para.add_run(new_text)
    if bold is not None:
        run.bold = bold
    run.font.name = font_name
    run.font.size = font_size


def _insert_paragraph_after(ref_para: Paragraph, text: str) -> Paragraph:
    """
    Insert a new paragraph immediately after ref_para, cloning its
    paragraph properties (pPr) and first run's formatting (rPr).

    Returns the new Paragraph so callers can chain insertions in order.
    """
    ref_p = ref_para._element

    # Clone paragraph properties (spacing, indent, borders, etc.)
    new_p = OxmlElement("w:p")
    p_pr = ref_p.find(qn("w:pPr"))
    if p_pr is not None:
        new_p.append(deepcopy(p_pr))

    # Clone run properties from the first run of the reference paragraph
    first_r = ref_p.find(qn("w:r"))
    r_pr_clone = None
    if first_r is not None:
        r_pr = first_r.find(qn("w:rPr"))
        if r_pr is not None:
            r_pr_clone = deepcopy(r_pr)

    # Build the run
    new_r = OxmlElement("w:r")
    if r_pr_clone is not None:
        new_r.append(r_pr_clone)
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_r.append(new_t)
    new_p.append(new_r)

    # Insert after the reference element
    ref_p.addnext(new_p)

    return Paragraph(new_p, ref_para._parent)


# ── Utilities ─────────────────────────────────────────────────────────────────


def _safe_filename(text: str) -> str:
    """Sanitize `text` for use as a filename component.

    Pure sanitizer — no hashing, no per-input discrimination. Applies:
      1. Strip filesystem-unsafe chars via `[^\\w\\s-]`.
      2. Collapse whitespace to `_`, then truncate to 50 chars.
      3. Fall back to the "unnamed" sentinel when the sanitized value
         is empty (whitespace-only input, Cf-only input the regex
         erased, etc.).

    Per-input collision discrimination is the caller's responsibility —
    see `_compose_output_base` for how the renderer combines the
    sanitizer output with a per-job-url discriminator. Making
    discrimination a caller concern keeps this helper reusable for
    future callers (log slugs, Notion page slugs, arbitrary user-
    input path segments) that do NOT want a mandatory `_XXXX` hash
    tail on English-language happy-path inputs.
    """
    safe = re.sub(r"[^\w\s-]", "", text)
    safe = re.sub(r"\s+", "_", safe.strip())
    return safe[:50] or "unnamed"


def _compose_output_base(job: dict, kind: str) -> str:
    """Compose the renderer's output filename base as
    ``{company}_{title}_{disc}_{kind}`` — where `disc` is a 4-hex
    per-job-url discriminator derived via ``blake2b(job['url'])``.

    Altitude rationale (PR #12 iter-2 pivot):
      - The parser's dedup contract (M7 + Phase 5 iter-2/iter-3) uses
        `(title, url)` as the fallback dedup key when
        `company == "Unknown"` (or empty-shape company). So two
        distinct URL-less-company jobs are guaranteed distinct URLs.
        The renderer path that most needs disambiguation — two same-
        title `company='Unknown'` jobs — is directly addressed by the
        URL-derived tail. For real-company jobs the parser's
        `(title, company)` dedup key already eliminates same-name
        collisions upstream, so the URL tail is defense-in-depth
        against non-parser code paths (test fixtures, future
        ingestion) that might not have run the same dedup.
      - The prior altitude (hash appended inside `_safe_filename`)
        used `text` (the company/title string) as the hash input,
        so two `company='Unknown'` jobs produced the same hash tail
        and silently overwrote each other. Moving the hash to the
        composite base with the URL as input eliminates that class.
      - Only ONE `_XXXX` tail lands on the base (not two — one per
        `_safe_filename` call), so multibyte inputs where per-
        component truncation could push the composite over
        `ENAMETOOLONG` see a single trailing discriminator instead
        of one per component.

    URL contract (iter-3 hardening): `_compose_output_base` REQUIRES
    a non-empty `job['url']` and raises `ValueError` if absent.
    Silent fallback (`.get('url') or ''`) would map every URL-less job
    to a fixed `blake2b(b'').hexdigest()[:4]` disc — so two same-
    company, same-title, URL-less jobs would still collide and
    silently overwrite each other. This is the exact M7 class the
    pivot targets, and any caller producing job dicts without a URL
    (test fixtures, upstream refactors bypassing the parser) needs
    to fail loud rather than degrade to a shared-path silent write.

    `errors='surrogatepass'` on `job['url'].encode(...)` handles the
    corner case where the URL string contains an isolated Unicode
    surrogate (BeautifulSoup can produce these under specific
    malformed-HTML paths). `blake2b` requires bytes, and default
    utf-8 encoding raises on a lone surrogate; `surrogatepass`
    round-trips them into the hash input so the discriminator remains
    computable rather than exploding at render time.
    """
    url = job.get("url")
    # Iter-3: check TYPE first so misleading "populate a URL" errors don't
    # fire for numeric-0 / False / bytes / Path callers whose actual bug
    # is the type, not the emptiness. TypeError names the wrong shape;
    # ValueError names the missing/whitespace content.
    if url is not None and not isinstance(url, str):
        raise TypeError(
            f"_compose_output_base requires job['url'] to be a str; got "
            f"{type(url).__name__}: {url!r}"
        )
    # Iter-3: `not url` catches empty string + None, but `url=' '` (or
    # any whitespace-only str) is truthy and would collapse to the fixed
    # `blake2b(b' ').hexdigest()[:4]` disc — same M7 class the guard was
    # written to prevent, just with a whitespace sentinel. Strip before
    # the emptiness check so whitespace-only URLs trip the ValueError arm
    # instead of silently colliding downstream.
    if not url or not url.strip():
        raise ValueError(
            "_compose_output_base requires a non-empty job['url']. Missing "
            "or whitespace-only URL would produce a fixed discriminator "
            "(blake2b of empty/whitespace bytes) that silently collides "
            "with every other URL-less job at the same (company, title) — "
            "the M7 class this helper was written to prevent. Callers "
            "constructing job dicts outside the parser must populate a "
            "URL identifier."
        )
    disc = blake2b(
        url.encode("utf-8", errors="surrogatepass"),
        digest_size=2,
    ).hexdigest()
    company = _safe_filename(job["company"])
    title = _safe_filename(job["title"])
    return f"{company}_{title}_{disc}_{kind}"
