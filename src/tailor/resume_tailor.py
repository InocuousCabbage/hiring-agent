"""
tailor/resume_tailor.py — Tailor a lane base resume to a specific JD.

Reads the base_resume.docx to extract current content, then calls Claude
to generate a tailored version. Returns a JSON-structured dict that the
renderer applies back to the template.

Resume structure (what can/cannot be changed):
  ✓ Tagline         — REWRITE per JD
  ✓ Summary         — REWRITE per JD
  ✓ KEY SKILLS      — REPLACE all 9 (3×3 table)
  ✓ Role 0 bullets  — EDIT (primary role)
  ✓ Role 1 bullets  — EDIT (secondary role)
  ✓ Role 2 bullets  — EDIT (tertiary role)
  ✓ Role 3 bullets  — SHORTEN (fourth role) — 1 line max
  ✗ Name / Contact  — DO NOT TOUCH
  ✗ Education       — DO NOT TOUCH
"""

import json
import re
from pathlib import Path

import structlog

from llm import call_claude

log = structlog.get_logger()

_MODEL = "claude-sonnet-4-6"

SYSTEM_PROMPT = """\
You are acting as both the hiring manager for this specific role AND an expert resume writer. \
Tailor the candidate's one-page marketing resume to make the hiring manager immediately want to \
schedule an interview.

MINDSET — think like the hiring manager for THIS exact role:
  • Before finalizing each bullet, ask: "Would this make me want to interview this person?"
  • If a bullet is generic enough to appear on anyone's resume, rewrite it with specific \
details from the project bank.
  • Integrate JD keywords naturally — never keyword-stuff or repeat mechanically.
  • The resume must look like a skilled human wrote it, not an AI.

RESUME STRUCTURE — what you can and cannot touch:
  • Name + Contact                                                                → DO NOT TOUCH
  • Tagline                                                                       → REWRITE
  • QUALIFICATION SUMMARY                                                         → REWRITE
  • KEY SKILLS (renderer adds "✓ " prefix — do not include it in your output)    → REPLACE all 9
  • Role 0 — Primary role (most recent full-time position)                       → EDIT bullets
  • Role 1 — Secondary role (consulting / side role)                             → EDIT bullets
  • Role 2 — Tertiary role (freelance / contract work)                           → EDIT bullets
  • Role 3 — Fourth role (earlier / less relevant position)                      → SHORTEN — compress to 1-2 lines maximum to free space for more relevant roles above.
  • Education                                                                     → DO NOT TOUCH

━━━ TAGLINE ━━━
• 3 key areas pulled directly from the JD's top priorities, separated by " | "
• Reads like a concise value proposition
• Example: "Product Marketing Strategy | Lifecycle Management | Market Analysis & Growth"

━━━ QUALIFICATION SUMMARY ━━━
• The summary MUST be under 300 characters. This is a hard limit. 2-3 SHORT sentences.
• NEVER first-person — do not use "I", "my", or "me"
• Direct, professional tone
• Must include: years of experience, industry context, strongest skills, and one top accomplishment
• Good example: "Product-oriented marketer with 2+ years driving go-to-market campaigns and \
CRM attribution. Track record of improving conversion rates through data-driven optimization \
across B2B and B2C markets."

━━━ KEY SKILLS — exactly 9 items ━━━
• 1–3 words per skill (e.g. "Marketing Automation", "Marketing Attribution", "A/B Testing")
• Hard skills ONLY — no soft skills (never "communication", "leadership", "teamwork")
• Drawn from skills the JD explicitly requires AND skills the candidate demonstrably has
• Do not list any skill the candidate cannot speak to from real experience

━━━ COMPANY ATTRIBUTION — NON-NEGOTIABLE ━━━
CRITICAL: Each role's bullets must ONLY describe work done at THAT specific company. \
Never place a project from one company under a different role.

  • Role 0 (Primary role): Where most project bank entries were done. \
Concentrate project-backed bullets in this role. Use 3–5 bullets drawing on the project bank.
  • Role 1 (Secondary role): A consulting or side role. \
Write exactly 1 bullet about client advisory work — channel audits, positioning, go-to-market \
recommendations. Do NOT reference projects from the primary role here.
  • Role 2 (Tertiary role): Freelance or contract work. Write exactly 1 bullet about \
independent work — site builds, analytics setup, client projects. \
Do NOT reference projects from the primary role here.
  • Role 3 (Fourth role): Compress to 1 short line — \
just enough to confirm the role existed. No embellishment.

━━━ BULLET POINTS — MOST CRITICAL SECTION ━━━
• Write ACCOMPLISHMENTS, not responsibilities
• Lead with the RESULT or IMPACT first, then explain how it was achieved
  ✓ GOOD: "Achieved a 20% cost savings on shipping by negotiating a new long-term deal with a key vendor"
  ✗ BAD:  "Responsible for negotiating vendor deals" — this is a job description, not an achievement
• Every bullet must include a MEASURABLE RESULT: a number, percentage, dollar figure, or concrete outcome
• Each bullet MUST be under 150 characters — this is a hard limit, no exceptions. Count carefully.
• Keep the TOTAL bullet count across all roles to 8–10 bullets maximum. Fewer, punchier bullets
  are better than more long ones. Do not pad roles with weak bullets to hit a minimum.
• No first-person ("I", "my") and no third-person ("he", "she")
• Start each bullet with a strong action verb

FORBIDDEN WORDS — never use any of these (dead giveaways of AI writing):
  meticulously, pioneered, prowess, realm, helm, spearheaded, orchestrated, leveraged, synergy,
  holistic, robust, cutting-edge, innovative, dynamic, streamlined, fostered, cultivated,
  garnered, harnessed, catapulted, propelled, instrumental, transformative, unparalleled, stellar

ALSO FORBIDDEN — em-dashes (—): never use em-dashes anywhere in the output. Use a comma,
  semicolon, or period instead. This includes bullets, the summary, and the tagline.

USE INSTEAD natural professional verbs:
  built, led, drove, created, launched, managed, improved, reduced, increased, developed,
  designed, implemented, tracked, analyzed, grew, delivered, generated, established, completed, resolved

━━━ NON-NEGOTIABLE RULES ━━━
1. TRUTH ONLY — every metric, tool, outcome, title, employer, and date must come from the
   project bank OR the current resume bullets. Never invent or fabricate anything.
2. 1 PAGE MAX — 8–10 bullets total across all roles; each bullet must be under 150 characters.
3. NO KEYWORD STUFFING — weave JD terms in naturally.
4. PROJECTS — reference 1–2 from the project bank using real metrics. No links or URLs.
5. GAPS — if the JD needs a skill the candidate lacks, translate the closest real experience
   to address it. Log the gap in gaps_noted; do NOT claim the skill outright.

━━━ CRITICAL — YOU MUST ALWAYS RETURN VALID JSON ━━━
Even if the job is a poor fit, you MUST return the JSON structure below. Never write prose
explanations instead of JSON. Use confidence_score to signal fit level:
  80–100  Strong match — directly relevant experience and skills
  50–79   Moderate match — transferable skills, some gaps
  0–49    Weak match — significant skill or industry mismatch
For low-confidence roles, populate every field as best you can using the closest transferable
experience. A confidence_score below the pipeline threshold will trigger a graceful skip —
but that decision belongs to the pipeline, not to you.

OUTPUT — return ONLY valid JSON, no markdown fences:
{
  "confidence_score": <integer 0–100>,
  "tagline": "<rewritten tagline>",
  "summary": "<rewritten summary, exactly 2–3 sentences, no first person>",
  "skills": ["Skill 1", "Skill 2", "Skill 3", "Skill 4", "Skill 5",
             "Skill 6", "Skill 7", "Skill 8", "Skill 9"],
  "roles": [
    {"index": 0, "bullets": ["bullet 1", "bullet 2", "bullet 3", "bullet 4"]},
    {"index": 1, "bullets": ["1 bullet — secondary/consulting role work only"]},
    {"index": 2, "bullets": ["1 bullet — freelance/contract work only"]},
    {"index": 3, "bullets": ["1 short line summarising fourth role"]}
  ],
  "gaps_noted": ["gap description", ...],
  "keywords_integrated": ["keyword1", "keyword2", ...]
}"""


def tailor_resume(
    jd_text: str,
    lane: dict,
    project_bank: list[dict],
    config: dict,
) -> dict:
    """
    Tailor the base resume content to match the JD.

    Returns a dict with keys: tagline, summary, skills, roles,
    gaps_noted, keywords_integrated, lane.
    """
    # Projects relevant to this lane
    relevant_projects = [
        p for p in project_bank
        if lane["name"] in p.get("lane", [])
    ]

    # Current resume text — gives Claude grounding before it rewrites
    template_path = Path(__file__).parent.parent.parent / lane.get(
        "template", "templates/resumes/base_resume.docx"
    )
    current_resume = _extract_resume_context(template_path)

    prompt = f"""Tailor this resume to the job description below.

<job_description>
{jd_text[:4000]}
</job_description>

<lane>
{lane['label']}
</lane>

<current_resume>
Tagline: {current_resume['tagline']}

Summary: {current_resume['summary']}

Skills: {', '.join(current_resume['skills'])}

Role 0 — {current_resume['roles'][0]['header']}
{_fmt_bullets(current_resume['roles'][0]['bullets'])}

Role 1 — {current_resume['roles'][1]['header']}
{_fmt_bullets(current_resume['roles'][1]['bullets'])}

Role 2 — {current_resume['roles'][2]['header']}
{_fmt_bullets(current_resume['roles'][2]['bullets'])}
</current_resume>

<project_bank>
{_format_project_bank(relevant_projects)}
</project_bank>

<constraints>
- Roles to edit: {config.get('max_roles_to_edit', 3)} (indices 0, 1, 2 only)
- Skills: exactly 9 items
- Must fit 1 page — keep bullets concise
- Reference 1–2 projects from the bank with real metrics
</constraints>

Return the tailored resume as JSON."""

    raw = call_claude(prompt, model=_MODEL, system=SYSTEM_PROMPT).strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)

    try:
        result = json.loads(raw)
    except json.JSONDecodeError as e:
        log.error("resume_tailor.json_parse_error", error=str(e), raw=raw[:300])
        result = {
            "confidence_score": 0,
            "tagline": "",
            "summary": "",
            "skills": [],
            "roles": [],
            "gaps_noted": ["Model returned prose instead of JSON — likely a very poor fit"],
            "keywords_integrated": [],
        }

    confidence = result.get("confidence_score", 100)
    log.info("resume_tailor.confidence", score=confidence)

    result = _strip_em_dashes(result)
    result["lane"] = lane
    return result


# ── Helpers ───────────────────────────────────────────────────────────────────

def _extract_resume_context(template_path: Path) -> dict:
    """
    Extract current tagline, summary, skills, and role bullets from the docx.
    Drives the prompt context so Claude can see what it's adapting.
    """
    from docx import Document

    doc = Document(template_path)
    paras = [p.text.strip() for p in doc.paragraphs]

    context: dict = {"tagline": "", "summary": "", "skills": [], "roles": []}

    # Tagline: first paragraph with "|" that isn't contact info
    for p in paras:
        if " | " in p and "@" not in p and "(" not in p:
            context["tagline"] = p
            break

    # Summary: paragraph immediately after the QUALIFICATION SUMMARY header
    for i, p in enumerate(paras):
        if p.upper() == "QUALIFICATION SUMMARY" and i + 1 < len(paras):
            context["summary"] = paras[i + 1]
            break

    # Skills from the first table (3×3 grid); strip the "✓ " prefix the renderer adds
    if doc.tables:
        context["skills"] = [
            cell.text.strip().lstrip("✓").strip()
            for row in doc.tables[0].rows
            for cell in row.cells
            if cell.text.strip()
        ]

    # Roles: find headers by date patterns; collect bullets until next header
    date_re = re.compile(
        r"(January|February|March|April|May|June|July|August|September|"
        r"October|November|December)\s+\d{4}"
    )
    header_indices = [i for i, p in enumerate(paras) if date_re.search(p)]

    for role_idx, (start, end) in enumerate(
        zip(header_indices, header_indices[1:] + [len(paras)])
    ):
        if role_idx >= 3:  # Only first 3 roles are editable
            break
        bullets = [
            p for p in paras[start + 1 : end]
            if p and not p.upper().startswith(("KEY SKILLS", "EDUCATION", "RELEVANT"))
        ]
        context["roles"].append({
            "index": role_idx,
            "header": paras[start],
            "bullets": bullets,
        })

    return context


def _clean_text(s: str) -> str:
    """Replace em/en-dashes then scrub the punctuation artifacts they leave behind."""
    s = s.replace("—", ", ").replace("–", "-")
    s = s.replace(" ,", ",")
    s = s.replace(" .", ".")
    s = re.sub(r"  +", " ", s)
    s = s.replace(",  ", ", ")
    return s.strip()


def _strip_em_dashes(result: dict) -> dict:
    """Clean em-dashes and resulting punctuation artifacts from all text fields."""
    result["tagline"] = _clean_text(result.get("tagline", ""))
    result["summary"] = _clean_text(result.get("summary", ""))
    result["skills"]  = [_clean_text(s) for s in result.get("skills", [])]
    for role in result.get("roles", []):
        role["bullets"] = [_clean_text(b) for b in role.get("bullets", [])]
    return result


def _fmt_bullets(bullets: list[str]) -> str:
    return "\n".join(f"  • {b}" for b in bullets) if bullets else "  (none)"


def _format_project_bank(projects: list[dict]) -> str:
    if not projects:
        return "(No projects available for this lane)"

    lines = []
    for p in projects:
        lines.append(
            f"[{p.get('id', '?')}] {p['name']} "
            f"({p.get('company', '')}, {p.get('date_range', '')})"
        )
        lines.append(f"  {p.get('summary', '').strip()}")
        for b in p.get("bullets", []):
            lines.append(f"  • {b}")
        for m in p.get("metrics", []):
            lines.append(f"  📊 {m}")
        lines.append(f"  Tools: {', '.join(p.get('tools_used', []))}")
        lines.append("")

    return "\n".join(lines)
