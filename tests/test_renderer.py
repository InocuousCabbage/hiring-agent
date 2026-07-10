#!/usr/bin/env python3
"""
tests/test_renderer.py — Render tailored resume + cover letter to DOCX/PDF.

Usage:
    python tests/test_renderer.py

Loads tailored JSON from tests/test_tailoring.py's saved run if available,
or uses a realistic hardcoded sample based on the actual resume structure.
Outputs both files to test_data/output/ and prints paths for visual inspection.
"""

import json
import sys
import zipfile
from datetime import date
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent / ".env")

from pdf_gen.renderer import render_resume, render_cover_letter

ROOT       = Path(__file__).parent.parent
OUTPUT_DIR = ROOT / "test_data" / "output"
SAVED_JD   = ROOT / "test_data" / "sample_jd.json"

# ── Sample data ───────────────────────────────────────────────────────────────
# Realistic example based on the actual resume + test_tailoring.py output.
# This is used when there's no live tailoring run to pull from.

SAMPLE_JOB = {
    "title": "Lead Product Marketing Manager",
    "company": "Group O",
    "location": "United States (Remote)",
    "salary": "$55–$62.5/hr",
    # PR #12 iter-3: `_compose_output_base` now REQUIRES a non-empty URL
    # to derive the per-job discriminator (parser-guaranteed unique per job).
    # Prior fixture had no `url` — silently invoked `blake2b(b'')` producing
    # a fixed discriminator across all renders. A caller test that rendered
    # two jobs with this fixture would collide on the same disk path.
    "url": "https://example.com/sample-alert/group-o-pmm",
}

SAMPLE_LANE = {
    "name": "pmm",
    "label": "Product Marketing (PMM)",
    "template": "templates/resumes/base_resume.docx",
}

SAMPLE_TAILORED_RESUME = {
    "tagline": "Product & Offer Implementation | Cross-Functional Leadership | Marketing Systems & Analytics",
    "summary": (
        "Marketing professional with 2+ years leading end-to-end implementation of "
        "marketing systems, CRM integrations, and multi-variable offer frameworks across "
        "B2C and B2B portfolios. Proven ability to translate complex business requirements "
        "into technical documentation, manage cross-functional initiatives from concept "
        "through launch, and drive measurable pipeline outcomes through structured "
        "program execution."
    ),
    "skills": [
        "Cross-Functional Program Management",
        "Marketing Requirements Documentation",
        "CRM & System Integration",
        "Product Lifecycle Management",
        "Campaign & Offer Execution",
        "Marketing Attribution Modeling",
        "Agile / Iterative Delivery",
        "Data Analysis & Reporting",
        "A/B Testing & CRO",
    ],
    "roles": [
        {
            "index": 0,
            "bullets": [
                "Led end-to-end implementation of marketing automation platform integrated "
                "with CRM, migrating 100,000+ records and generating "
                "$1M+ in pipeline to date.",
                "Designed a custom CRM integration strategy, architecting field-level data "
                "flow across Contacts, Leads, Opportunities, and Deals, translating complex "
                "business logic into precise technical requirements.",
                "Built and launched a marketing attribution infrastructure defining "
                "first-touch, campaign, referral, and funnel-stage fields across 2 business "
                "lines, enabling cross-channel performance reporting.",
                "Drove 223% average improvement in funnel conversion rates by leading CRO "
                "initiatives including A/B testing, behavioral analysis, and landing page "
                "optimization.",
                "Collaborated across Sales, IT, and Operations teams to implement revenue "
                "attribution models and deliver regular leadership readouts via BI "
                "dashboards.",
            ],
        },
        {
            "index": 1,
            "bullets": [
                "Led a regional product marketing strategy sprint for a new market entry, "
                "defining audience segmentation, channel prioritization, and go-to-market "
                "initiative sequencing.",
                "Advise clients on offer positioning, funnel architecture, and campaign "
                "execution, translating business objectives into actionable marketing and "
                "system requirements.",
                "Audit analytics infrastructure and channel performance to surface "
                "optimization opportunities, delivering structured recommendations aligned "
                "to client KPIs.",
            ],
        },
        {
            "index": 2,
            "bullets": [
                "Design and launch SEO-ready websites with integrated analytics, automation, "
                "and UX improvements; manage ongoing performance, accessibility, and content "
                "updates for multiple clients.",
            ],
        },
    ],
    "gaps_noted": [
        "No direct DIRECTV/telecom billing system experience (C3, STMS, Amdocs); "
        "analogous CRM integration experience with marketing automation and CRM platforms.",
    ],
    "keywords_integrated": [
        "product marketing", "offer implementation", "cross-functional leadership",
        "marketing requirements documentation", "go-to-market", "attribution",
    ],
    "lane": SAMPLE_LANE,
}

SAMPLE_COVER_LETTER = {
    "paragraphs": [
        "Telecom offer implementation is one of the more technically demanding corners of "
        "product marketing — managing hundreds of offer variables, coordinating across "
        "billing, engineering, and GTM teams, and keeping four to six initiatives moving "
        "simultaneously requires the kind of cross-functional discipline that most marketing "
        "roles never develop. That operational depth is where most of my recent work has lived.",

        "At Acme Corp, I led the end-to-end architecture, testing, and launch of a "
        "marketing automation and CRM integration, translating business requirements "
        "into precise technical specs for engineering and systems partners, managing a custom "
        "sync model across leads, opportunities, and project records, and migrating "
        "100,000+ customer records without disrupting the active sales pipeline. The project "
        "generated $1M+ in pipeline to date. That work maps directly to "
        "what this role requires: breaking down marketing and product needs into actionable "
        "technical requirements and managing implementation across multiple systems and "
        "stakeholders.",

        "I also led a marketing attribution infrastructure build in the CRM, scoping field "
        "logic changes, coordinating development across two business lines, and extending "
        "attribution continuity across lead, opportunity, and project records to enable "
        "reporting by channel, campaign, category, and funnel stage.",

        "I'd welcome a conversation about how this experience applies to the DIRECTV product "
        "and offer implementation work at Group O.",
    ],
    "projects_referenced": ["proj_001", "proj_002"],
}


# ── Main ──────────────────────────────────────────────────────────────────────


def main() -> None:
    # Try to load live job metadata if the JD fetcher saved it
    job = SAMPLE_JOB
    if SAVED_JD.exists():
        try:
            saved = json.loads(SAVED_JD.read_text())
            job = saved.get("job", SAMPLE_JOB)
            print(f"Loaded job metadata from {SAVED_JD}")
        except Exception:
            pass

    today = date.today().isoformat()
    output_dir = OUTPUT_DIR / today

    print(f"\nRendering for: {job['title']} @ {job['company']}")
    print(f"Output dir   : {output_dir}")
    print(f"Date         : {today}")

    # ── Resume ───────────────────────────────────────────────────────────────
    print("\n[1/2] Rendering resume DOCX → PDF...")
    resume_pdf, resume_docx = render_resume(
        tailored_resume=SAMPLE_TAILORED_RESUME,
        lane=SAMPLE_LANE,
        job=job,
        date_str=today,
        output_dir=output_dir,
    )
    if resume_pdf is None:
        print(f"      PDF  : (no PDF converter installed)")
    else:
        print(f"      PDF  : {resume_pdf}")
    print(f"      DOCX : {resume_docx}")

    # ── Cover letter ─────────────────────────────────────────────────────────
    print("\n[2/2] Rendering cover letter DOCX → PDF...")
    cl_pdf, cl_docx = render_cover_letter(
        cover_letter=SAMPLE_COVER_LETTER,
        job=job,
        date_str=today,
        output_dir=output_dir,
    )
    if cl_pdf is None:
        print(f"      PDF  : (no PDF converter installed)")
    else:
        print(f"      PDF  : {cl_pdf}")
    print(f"      DOCX : {cl_docx}")

    # ── Verify files exist and have content ──────────────────────────────────
    print("\n── File check ──")
    ok = True
    file_checks = [
        ("Resume DOCX",  resume_docx),
        ("Cover DOCX",   cl_docx),
    ]
    if resume_pdf is not None:
        file_checks.insert(0, ("Resume PDF", resume_pdf))
    if cl_pdf is not None:
        file_checks.insert(-1, ("Cover PDF", cl_pdf))
    for label, path in file_checks:
        if path.exists():
            size = path.stat().st_size
            status = "✓ OK" if size > 1000 else "⚠ suspiciously small"
            print(f"  {label:15s}: {size:>8,} bytes  {status}  {path.name}")
        else:
            print(f"  {label:15s}: ✗ NOT FOUND at {path}")
            ok = False

    if ok:
        print("\nOpen to inspect:")
        if resume_pdf is not None:
            print(f"  open '{resume_pdf}'")
        if cl_pdf is not None:
            print(f"  open '{cl_pdf}'")
    else:
        print("\n⚠ One or more files missing — check logs above.")


# ── Pytest unit tests ─────────────────────────────────────────────────────────

_TEMPLATE_PATH = ROOT / "templates" / "resumes" / "base_resume.docx"
_TEMPLATE_MISSING = not _TEMPLATE_PATH.exists()
_TEMPLATE_SKIP_REASON = (
    f"Base resume template missing at {_TEMPLATE_PATH} "
    "(user-supplied + gitignored — set up per SETUP.md to run renderer tests)."
)


@pytest.mark.skipif(_TEMPLATE_MISSING, reason=_TEMPLATE_SKIP_REASON)
def test_render_resume_returns_pdf_and_docx_tuple(tmp_path):
    """render_resume returns (Optional[pdf_path], docx_path); DOCX is a valid OOXML file."""
    pdf_path, docx_path = render_resume(
        tailored_resume=SAMPLE_TAILORED_RESUME,
        lane=SAMPLE_LANE,
        job=SAMPLE_JOB,
        date_str=date.today().isoformat(),
        output_dir=tmp_path,
    )

    # DOCX is always produced — even if PDF conversion fails
    assert docx_path.exists(), f"DOCX not written at {docx_path}"
    assert docx_path.suffix == ".docx"
    assert zipfile.is_zipfile(docx_path), "DOCX must be a valid OOXML/ZIP file"

    doc = Document(str(docx_path))
    # base_resume.docx has 26+ paragraphs (per renderer.py docstring)
    assert len(doc.paragraphs) > 20, (
        f"Expected >20 paragraphs in tailored resume DOCX, got {len(doc.paragraphs)}"
    )

    # PDF is Optional: an extant Path when LibreOffice/docx2pdf is available,
    # None when the renderer fell back to DOCX-only.
    assert pdf_path is None or pdf_path.exists(), (
        f"PDF slot must be None or an extant Path, got {pdf_path}"
    )


def test_render_cover_letter_returns_pdf_and_docx_tuple(tmp_path):
    """render_cover_letter returns (Optional[pdf_path], docx_path); DOCX is valid.

    Cover letters are built from scratch — no resume template required, so this
    test does not need the skip guard that the resume test has.
    """
    pdf_path, docx_path = render_cover_letter(
        cover_letter=SAMPLE_COVER_LETTER,
        job=SAMPLE_JOB,
        date_str=date.today().isoformat(),
        output_dir=tmp_path,
    )

    assert docx_path.exists(), f"Cover letter DOCX not written at {docx_path}"
    assert docx_path.suffix == ".docx"
    assert zipfile.is_zipfile(docx_path)

    doc = Document(str(docx_path))
    # Cover letter has applicant name + contact + N paragraphs from SAMPLE_COVER_LETTER
    assert len(doc.paragraphs) >= len(SAMPLE_COVER_LETTER["paragraphs"]) + 2

    assert pdf_path is None or pdf_path.exists()


def test_render_resume_returns_none_pdf_on_fallback(tmp_path, _no_pdf_converter):
    """HIGH-3 (structural fix): When no PDF converter is available, render_resume
    must return (None, docx_path) — NOT (docx_path, docx_path).

    Returning the same Path twice loses the ops signal that PDF conversion failed.
    None in the PDF slot lets callers detect fallback unambiguously.

    iter-3 (finding #5): stub extraction. The inline `monkeypatch.setattr(
    renderer_mod, "_find_libreoffice", ...) + sys.platform = "linux"` block
    duplicated the `_no_pdf_converter` fixture defined below — a drift
    hazard if the platform check moves off `sys.platform`. Both tests
    now go through the fixture so a single edit keeps them in sync.
    """
    if _TEMPLATE_MISSING:
        pytest.skip(_TEMPLATE_SKIP_REASON)

    pdf_path, docx_path = render_resume(
        tailored_resume=SAMPLE_TAILORED_RESUME,
        lane=SAMPLE_LANE,
        job=SAMPLE_JOB,
        date_str=date.today().isoformat(),
        output_dir=tmp_path,
    )

    assert pdf_path is None, (
        f"Expected pdf_path = None on fallback (no converter), got {pdf_path}. "
        "Returning the same path twice loses the ops signal."
    )
    assert docx_path is not None and docx_path.exists()
    assert docx_path.suffix == ".docx"


def test_render_cover_letter_returns_none_pdf_on_fallback(tmp_path, _no_pdf_converter):
    """HIGH-3: cover letter renderer must return (None, docx_path) on fallback.

    iter-3 (finding #5): uses _no_pdf_converter fixture — see paired
    test_render_resume_returns_none_pdf_on_fallback for the rationale."""
    pdf_path, docx_path = render_cover_letter(
        cover_letter=SAMPLE_COVER_LETTER,
        job=SAMPLE_JOB,
        date_str=date.today().isoformat(),
        output_dir=tmp_path,
    )

    assert pdf_path is None, (
        f"Expected pdf_path = None on fallback, got {pdf_path}"
    )
    assert docx_path is not None and docx_path.exists()
    assert docx_path.suffix == ".docx"


# ── M23: renderer content assertions ──────────────────────────────────────────
# Prior state: the renderer tests asserted only zip validity + paragraph
# count > 20 — an untailored base_resume.docx (all paragraphs untouched)
# passed. If _fill_resume_template silently no-op'd (broken .runs walk, empty
# section handler, etc.), the shipped DOCX would be the un-tailored template
# and the tests would never notice.
#
# These tests read the DOCX text back and assert the tailored summary +
# bullet content is actually present. Mutation check: no-op the placeholder
# fill loop in renderer._fill_resume_template — the summary text vanishes
# from the DOCX and this test fails.


def _docx_full_text(docx_path: Path) -> str:
    """Extract concatenated paragraph text from a DOCX (including tables)."""
    doc = Document(str(docx_path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


@pytest.fixture
def _no_pdf_converter(monkeypatch):
    """Phase 5 iter-2 (finding #12): stub the LibreOffice + docx2pdf paths
    so renderer tests that only assert DOCX content don't spawn a real
    LibreOffice subprocess (which can take up to 60s per call — see
    _find_libreoffice's subprocess.TimeoutExpired handler). Same pattern
    the pre-existing `test_render_resume_returns_none_pdf_on_fallback`
    test uses; applied here to the M21 + M23 tests that call render_*
    without stubbing.
    """
    import pdf_gen.renderer as renderer_mod
    monkeypatch.setattr(renderer_mod, "_find_libreoffice", lambda: None)
    # Non-darwin platform skips the docx2pdf branch entirely.
    monkeypatch.setattr(sys, "platform", "linux")


@pytest.mark.skipif(_TEMPLATE_MISSING, reason=_TEMPLATE_SKIP_REASON)
def test_render_resume_docx_contains_tailored_summary_and_bullets(tmp_path, _no_pdf_converter):
    """M23: the rendered DOCX must contain the tailored summary + a
    signature phrase from at least one bullet. If the fill loop no-ops,
    the DOCX ships the base template and this test fails."""
    _, docx_path = render_resume(
        tailored_resume=SAMPLE_TAILORED_RESUME,
        lane=SAMPLE_LANE,
        job=SAMPLE_JOB,
        date_str=date.today().isoformat(),
        output_dir=tmp_path,
    )
    text = _docx_full_text(docx_path)

    # Tailored summary must be present verbatim (or nearly so — the fill
    # process may reflow whitespace, so we check a distinctive substring).
    summary_marker = "translate complex business requirements"
    assert summary_marker in text, (
        f"Tailored summary substring {summary_marker!r} missing from rendered "
        f"DOCX. Renderer likely no-op'd the summary fill.\n"
        f"First 400 chars of DOCX text:\n{text[:400]}"
    )

    # At least one distinctive bullet phrase from SAMPLE_TAILORED_RESUME
    # must be present.
    bullet_marker = "223% average improvement in funnel conversion"
    assert bullet_marker in text, (
        f"Tailored bullet substring {bullet_marker!r} missing from rendered "
        f"DOCX. Renderer likely dropped role bullets."
    )


def test_render_cover_letter_docx_contains_tailored_paragraphs(tmp_path, _no_pdf_converter):
    """M23: cover letter DOCX must contain the tailored paragraph body,
    not just the applicant name header. If _create_cover_letter_docx
    silently dropped the paragraphs, this test fails."""
    _, docx_path = render_cover_letter(
        cover_letter=SAMPLE_COVER_LETTER,
        job=SAMPLE_JOB,
        date_str=date.today().isoformat(),
        output_dir=tmp_path,
    )
    text = _docx_full_text(docx_path)

    # A distinctive phrase from the first SAMPLE_COVER_LETTER paragraph.
    marker = "Telecom offer implementation is one of the more technically demanding"
    assert marker in text, (
        f"Cover-letter tailored paragraph missing from rendered DOCX. "
        f"First 400 chars:\n{text[:400]}"
    )
    # And a second-paragraph phrase to prove multi-paragraph rendering.
    marker2 = "generated $1M+ in pipeline"
    assert marker2 in text, (
        f"Second cover-letter paragraph missing. First 400 chars:\n{text[:400]}"
    )


# ── M21: chdir render test ────────────────────────────────────────────────────
# Prior guard (`test_render_resume_pdf_uses_root_based_path` in
# test_review_fixes.py) is a source-text grep — it lints for the _ROOT/
# pattern but does NOT verify runtime behavior. A CWD-relative regression
# under a different spelling (e.g. `os.path.join(os.getcwd(), lane[...])`)
# would still ship.
#
# This test chdirs into a scratch dir, then invokes render_resume — a
# renderer that reads its template relative to CWD (not _ROOT) crashes with
# FileNotFoundError. Passes today because renderer resolves against _ROOT.


@pytest.mark.skipif(_TEMPLATE_MISSING, reason=_TEMPLATE_SKIP_REASON)
def test_render_resume_finds_template_from_arbitrary_cwd(tmp_path, monkeypatch, _no_pdf_converter):
    """M21: render_resume must resolve its template against the repo root,
    not the process CWD. Chdir into a scratch dir with no template beneath
    it and verify the render still succeeds.

    Mutation check: revert renderer.py's template_path to
    `Path(lane['template'])`. This test fails with FileNotFoundError,
    proving the guard is behavioral rather than source-grep."""
    # A scratch dir with NO templates/ subtree.
    scratch = tmp_path / "scratch_cwd"
    scratch.mkdir()
    output_dir = tmp_path / "out"

    monkeypatch.chdir(scratch)

    pdf_path, docx_path = render_resume(
        tailored_resume=SAMPLE_TAILORED_RESUME,
        lane=SAMPLE_LANE,
        job=SAMPLE_JOB,
        date_str=date.today().isoformat(),
        output_dir=output_dir,
    )
    # DOCX must exist under output_dir; renderer must NOT have raised
    # FileNotFoundError on the template.
    assert docx_path.exists(), (
        f"DOCX not rendered from arbitrary CWD — template likely resolved "
        f"against CWD instead of _ROOT.\n"
        f"CWD was: {scratch}\n"
        f"Expected output: {docx_path}"
    )
    assert docx_path.suffix == ".docx"


if __name__ == "__main__":
    main()
